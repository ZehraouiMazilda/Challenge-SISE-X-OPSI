"""
Firewall Log Analysis Report Generator
Section 1.1 - SISE-OPSIE Project 2026
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

# Set style for plots
plt.style.use('seaborn-v0_8-darkgrid')
sns.set_palette("husl")

def load_firewall_logs(filepath):
    """
    Load and preprocess firewall logs
    """
    # Column names based on the data structure
    columns = ['date', 'ip_source', 'ip_destination', 'protocol', 
               'source_port', 'dest_port', 'rule_id', 'action', 
               'interface_in', 'interface_out', 'firewall_num']
    
    # Load data
    df = pd.read_csv(filepath, 
                     sep=';',
                     names=columns,
                     parse_dates=['date'])
    
    # Filter: November 2025 to February 2026
    # Changed 02-29 to 02-28 (2026 is not a leap year)
    df = df[(df['date'] >= pd.Timestamp('2025-11-01')) & (df['date'] <= pd.Timestamp('2026-02-28'))]
    
    # Remove firewall number column (FW=6)
    df = df.drop('firewall_num', axis=1)
    
    print(f" Loaded {len(df):,} log entries")
    print(f" Date range: {df['date'].min()} to {df['date'].max()}")
    
    return df

def chapter1_rule_ranking(df):
    rule_counts = df['rule_id'].value_counts().reset_index()
    rule_counts.columns = ['Rule ID', 'Count']
    rule_counts['Percentage'] = (rule_counts['Count'] / len(df) * 100).round(2)
    
    # Create visualization
    fig, ax = plt.subplots(figsize=(10, 6))
    top_rules = rule_counts.head(15)
    
    bars = ax.barh(top_rules['Rule ID'].astype(str), top_rules['Count'])
    ax.set_xlabel('Number of Occurrences', fontsize=12)
    ax.set_ylabel('Rule ID', fontsize=12)
    ax.set_title('Top 15 Most Used Firewall Rules', fontsize=14, fontweight='bold')
    ax.invert_yaxis()
    
    # Add value labels
    for i, (count, pct) in enumerate(zip(top_rules['Count'], top_rules['Percentage'])):
        ax.text(count, i, f' {count:,} ({pct}%)', va='center', fontsize=9)
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return rule_counts, img_buffer

def chapter2_protocol_histogram(df):
    protocol_counts = df['protocol'].value_counts().reset_index()
    protocol_counts.columns = ['Protocol', 'Count']
    protocol_counts['Percentage'] = (protocol_counts['Count'] / len(df) * 100).round(2)
    
    # Create visualization
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = sns.color_palette("Set2", len(protocol_counts))
    
    bars = ax.bar(protocol_counts['Protocol'], protocol_counts['Count'], color=colors)
    ax.set_xlabel('Protocol', fontsize=12)
    ax.set_ylabel('Number of Flows', fontsize=12)
    ax.set_title('Protocol Usage Distribution', fontsize=14, fontweight='bold')
    
    # Add value labels
    for bar, count, pct in zip(bars, protocol_counts['Count'], protocol_counts['Percentage']):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{count:,}\n({pct}%)',
                ha='center', va='bottom', fontsize=10)
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return protocol_counts, img_buffer

def chapter3_top10_udp_rules(df):
    udp_df = df[df['protocol'] == 'UDP']
    
    if len(udp_df) == 0:
        return None, None
    
    top_udp_rules = udp_df['rule_id'].value_counts().head(10).reset_index()
    top_udp_rules.columns = ['Rule ID', 'Count']
    top_udp_rules['Percentage'] = (top_udp_rules['Count'] / len(udp_df) * 100).round(2)
    
    # Create visualization
    fig, ax = plt.subplots(figsize=(10, 6))
    
    bars = ax.barh(top_udp_rules['Rule ID'].astype(str), top_udp_rules['Count'], color='#2ecc71')
    ax.set_xlabel('Number of Occurrences', fontsize=12)
    ax.set_ylabel('Rule ID', fontsize=12)
    ax.set_title('Top 10 Most Used Rules - UDP Protocol', fontsize=14, fontweight='bold')
    ax.invert_yaxis()
    
    # Add value labels
    for i, (count, pct) in enumerate(zip(top_udp_rules['Count'], top_udp_rules['Percentage'])):
        ax.text(count, i, f' {count:,} ({pct}%)', va='center', fontsize=9)
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return top_udp_rules, img_buffer

def chapter4_top5_tcp_rules(df):
    tcp_df = df[df['protocol'] == 'TCP']
    
    top_tcp_rules = tcp_df['rule_id'].value_counts().head(5).reset_index()
    top_tcp_rules.columns = ['Rule ID', 'Count']
    top_tcp_rules['Percentage'] = (top_tcp_rules['Count'] / len(tcp_df) * 100).round(2)
    
    # Create visualization
    fig, ax = plt.subplots(figsize=(10, 6))
    
    bars = ax.barh(top_tcp_rules['Rule ID'].astype(str), top_tcp_rules['Count'], color='#3498db')
    ax.set_xlabel('Number of Occurrences', fontsize=12)
    ax.set_ylabel('Rule ID', fontsize=12)
    ax.set_title('Top 5 Most Used Rules - TCP Protocol', fontsize=14, fontweight='bold')
    ax.invert_yaxis()
    
    # Add value labels
    for i, (count, pct) in enumerate(zip(top_tcp_rules['Count'], top_tcp_rules['Percentage'])):
        ax.text(count, i, f' {count:,} ({pct}%)', va='center', fontsize=9)
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return top_tcp_rules, img_buffer

def chapter5_tcp_rules_ports_actions(df):
    """
    Chapter 5: Relationship between rules, destination ports, and actions (TCP only)
    """
    tcp_df = df[df['protocol'] == 'TCP']
    
    # Analysis by rule, port, and action
    analysis = tcp_df.groupby(['rule_id', 'dest_port', 'action']).size().reset_index(name='count')
    analysis = analysis.sort_values('count', ascending=False).head(20)
    
    # Create pivot for heatmap of top rules vs top ports
    top_rules = tcp_df['rule_id'].value_counts().head(10).index
    top_ports = tcp_df['dest_port'].value_counts().head(10).index
    
    heatmap_data = tcp_df[tcp_df['rule_id'].isin(top_rules) & tcp_df['dest_port'].isin(top_ports)]
    pivot = pd.crosstab(heatmap_data['rule_id'], heatmap_data['dest_port'])
    
    # Create visualization
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    
    # Heatmap
    sns.heatmap(pivot, annot=True, fmt='d', cmap='YlOrRd', ax=ax1, cbar_kws={'label': 'Count'})
    ax1.set_title('TCP Traffic Heatmap: Rules vs Destination Ports', fontsize=12, fontweight='bold')
    ax1.set_xlabel('Destination Port', fontsize=11)
    ax1.set_ylabel('Rule ID', fontsize=11)
    
    # Action distribution by top rules
    top_rules_action = tcp_df[tcp_df['rule_id'].isin(top_rules)].groupby(['rule_id', 'action']).size().unstack(fill_value=0)
    top_rules_action.plot(kind='bar', stacked=True, ax=ax2, color=['#e74c3c', '#2ecc71'])
    ax2.set_title('Action Distribution by Top Rules (TCP)', fontsize=12, fontweight='bold')
    ax2.set_xlabel('Rule ID', fontsize=11)
    ax2.set_ylabel('Count', fontsize=11)
    ax2.legend(title='Action', loc='upper right')
    ax2.tick_params(axis='x', rotation=45)
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return analysis, img_buffer

def chapter6_additional_security_analysis(df):
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
    
    # 1. Permit vs Deny over time
    df['date_only'] = df['date'].dt.date
    daily_actions = df.groupby(['date_only', 'action']).size().unstack(fill_value=0)
    daily_actions.plot(ax=ax1, color=['#e74c3c', '#2ecc71'], linewidth=2)
    ax1.set_title('Daily Traffic: Permit vs Deny', fontsize=12, fontweight='bold')
    ax1.set_xlabel('Date', fontsize=11)
    ax1.set_ylabel('Number of Flows', fontsize=11)
    ax1.legend(title='Action')
    ax1.grid(True, alpha=0.3)
    
    # 2. Top 10 Source IPs with denied traffic
    denied_ips = df[df['action'] == 'DENY']['ip_source'].value_counts().head(10)
    denied_ips.plot(kind='barh', ax=ax2, color='#e74c3c')
    ax2.set_title('Top 10 Source IPs with Most Denied Traffic', fontsize=12, fontweight='bold')
    ax2.set_xlabel('Number of Denied Flows', fontsize=11)
    ax2.set_ylabel('Source IP', fontsize=11)
    
    # 3. Traffic by hour of day
    df['hour'] = df['date'].dt.hour
    hourly_traffic = df.groupby('hour').size()
    hourly_traffic.plot(kind='bar', ax=ax3, color='#3498db')
    ax3.set_title('Traffic Distribution by Hour of Day', fontsize=12, fontweight='bold')
    ax3.set_xlabel('Hour', fontsize=11)
    ax3.set_ylabel('Number of Flows', fontsize=11)
    ax3.tick_params(axis='x', rotation=0)
    
    # 4. Top 15 destination ports
    top_ports = df['dest_port'].value_counts().head(15)
    top_ports.plot(kind='barh', ax=ax4, color='#9b59b6')
    ax4.set_title('Top 15 Most Targeted Destination Ports', fontsize=12, fontweight='bold')
    ax4.set_xlabel('Number of Flows', fontsize=11)
    ax4.set_ylabel('Destination Port', fontsize=11)
    ax4.invert_yaxis()
    
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer

def create_report(df, output_filename='Rapport_Analyse_Firewall.docx'):
    """
    Here we generate complete Word document report
    """
    doc = Document()
    
    # Title page
    title = doc.add_heading('Analyse des Logs Firewall IPTables', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Projet SISE-OPSIE 2026')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    date_para = doc.add_paragraph(f'Date du rapport: {datetime.now().strftime("%d/%m/%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Summary statistics
    doc.add_heading('Résumé Exécutif', 1)
    doc.add_paragraph(f"Nombre total de flux analysés: {len(df):,}")
    doc.add_paragraph(f"Période analysée: {df['date'].min().strftime('%d/%m/%Y')} - {df['date'].max().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Nombre d'IP sources uniques: {df['ip_source'].nunique():,}")
    doc.add_paragraph(f"Nombre de règles distinctes: {df['rule_id'].nunique()}")
    
    permit_count = len(df[df['action'] == 'PERMIT'])
    deny_count = len(df[df['action'] == 'DENY'])
    doc.add_paragraph(f"Flux autorisés: {permit_count:,} ({permit_count/len(df)*100:.1f}%)")
    doc.add_paragraph(f"Flux bloqués: {deny_count:,} ({deny_count/len(df)*100:.1f}%)")
    
    doc.add_page_break()
    
    # Chapter 1
    print(" Generating Chapter 1: Rule Ranking...")
    doc.add_heading('1. Classement des Règles les Plus Utilisées', 1)
    doc.add_paragraph(
        "Cette section présente le classement des règles de filtrage par ordre d'utilisation. "
        "L'identification des règles les plus sollicitées permet d'optimiser les performances "
        "du firewall en positionnant ces règles en début de chaîne."
    )
    rule_counts, img1 = chapter1_rule_ranking(df)
    doc.add_picture(img1, width=Inches(6))
    
    # Add top 10 rules table
    doc.add_paragraph()
    doc.add_paragraph("Top 10 des règles:", style='Intense Quote')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule ID'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Pourcentage'
    
    for _, row in rule_counts.head(10).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Rule ID'])
        row_cells[1].text = f"{row['Count']:,}"
        row_cells[2].text = f"{row['Percentage']}%"
    
    doc.add_page_break()
    
    # Chapter 2
    print(" Generating Chapter 2: Protocol Histogram...")
    doc.add_heading('2. Utilisation des Différents Protocoles', 1)
    doc.add_paragraph(
        "Cette analyse présente la distribution des protocoles réseau observés dans les logs. "
        "TCP et UDP représentent généralement la majorité du trafic."
    )
    protocol_counts, img2 = chapter2_protocol_histogram(df)
    doc.add_picture(img2, width=Inches(6))
    
    doc.add_paragraph()
    doc.add_paragraph("Statistiques par protocole:", style='Intense Quote')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Protocole'
    hdr_cells[1].text = 'Nombre de flux'
    hdr_cells[2].text = 'Pourcentage'
    
    for _, row in protocol_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Protocol']
        row_cells[1].text = f"{row['Count']:,}"
        row_cells[2].text = f"{row['Percentage']}%"
    
    doc.add_page_break()
    
    # Chapter 3
    print(" Generating Chapter 3: Top 10 UDP Rules...")
    doc.add_heading('3. Top 10 des Règles UDP', 1)
    doc.add_paragraph(
        "Analyse des règles les plus utilisées pour le protocole UDP. "
        "UDP est souvent utilisé pour les services DNS, streaming, et jeux en ligne."
    )
    top_udp_rules, img3 = chapter3_top10_udp_rules(df)
    
    if img3:
        doc.add_picture(img3, width=Inches(6))
        
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Rule ID'
        hdr_cells[1].text = 'Nombre'
        hdr_cells[2].text = 'Pourcentage'
        
        for _, row in top_udp_rules.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Rule ID'])
            row_cells[1].text = f"{row['Count']:,}"
            row_cells[2].text = f"{row['Percentage']}%"
    else:
        doc.add_paragraph(" Aucun trafic UDP détecté dans les logs analysés.")
    
    doc.add_page_break()
    
    # Chapter 4
    print(" Generating Chapter 4: Top 5 TCP Rules...")
    doc.add_heading('4. Top 5 des Règles TCP', 1)
    doc.add_paragraph(
        "Analyse des règles les plus utilisées pour le protocole TCP. "
        "TCP est le protocole principal pour les services web, SSH, FTP, et bases de données."
    )
    top_tcp_rules, img4 = chapter4_top5_tcp_rules(df)
    doc.add_picture(img4, width=Inches(6))
    
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rule ID'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Pourcentage'
    
    for _, row in top_tcp_rules.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Rule ID'])
        row_cells[1].text = f"{row['Count']:,}"
        row_cells[2].text = f"{row['Percentage']}%"
    
    doc.add_page_break()
    
    # Chapter 5
    print(" Generating Chapter 5: TCP Rules-Ports-Actions Analysis...")
    doc.add_heading('5. Rapprochement Règles-Ports-Actions (TCP)', 1)
    doc.add_paragraph(
        "Cette analyse croise les règles de filtrage avec les ports de destination et les actions "
        "pour le protocole TCP. Elle permet d'identifier les patterns de trafic et les règles "
        "de sécurité les plus sollicitées."
    )
    analysis, img5 = chapter5_tcp_rules_ports_actions(df)
    doc.add_picture(img5, width=Inches(6.5))
    
    doc.add_page_break()
    
    # Chapter 6
    print(" Generating Chapter 6: Additional Security Analysis...")
    doc.add_heading('6. Analyses Complémentaires de Sécurité', 1)
    doc.add_paragraph(
        "Cette section présente des analyses supplémentaires pertinentes pour la sécurité du SI: "
        "évolution temporelle du trafic, identification des sources malveillantes potentielles, "
        "patterns horaires, et ports les plus ciblés."
    )
    img6 = chapter6_additional_security_analysis(df)
    doc.add_picture(img6, width=Inches(6.5))
    
    doc.add_paragraph()
    doc.add_heading('Observations:', 2)
    doc.add_paragraph(
        "• L'évolution temporelle permet d'identifier des pics d'activité suspects\n"
        "• Les IP sources avec le plus de trafic bloqué sont potentiellement malveillantes\n"
        "• Les patterns horaires révèlent les périodes de forte activité\n"
        "• Les ports les plus ciblés indiquent les services exposés"
    )
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        f"Cette analyse des logs firewall IPTables sur la période de novembre 2025 à février 2026 "
        f"a permis d'examiner {len(df):,} flux réseau. Les principales observations sont:"
    )
    
    doc.add_paragraph(
        f"• La règle {rule_counts.iloc[0]['Rule ID']} est la plus utilisée "
        f"({rule_counts.iloc[0]['Percentage']}% du trafic total)"
    )
    
    doc.add_paragraph(
        f"• Le protocole {protocol_counts.iloc[0]['Protocol']} domine avec "
        f"{protocol_counts.iloc[0]['Percentage']}% du trafic"
    )
    
    doc.add_paragraph(
        f"• {deny_count/len(df)*100:.1f}% du trafic a été bloqué, indiquant une activité "
        f"de sécurité significative"
    )
    
    doc.add_paragraph(
        "\nRecommandations:"
    )
    doc.add_paragraph(
        "• Optimiser l'ordre des règles en plaçant les plus utilisées en début de chaîne\n"
        "• Investiguer les IP sources avec un taux élevé de trafic bloqué\n"
        "• Réviser les règles peu utilisées pour simplifier la configuration\n"
        "• Monitorer les pics d'activité suspects identifiés dans l'analyse temporelle"
    )
    
    # Save document
    doc.save(output_filename)
    print(f"\n Rapport généré avec succès: {output_filename}")

def main():
    """
    Main execution function
    """
    print("=" * 60)
    print("ANALYSE DES LOGS FIREWALL IPTABLES")
    print("Projet SISE-OPSIE 2026 - Section 1.1")
    print("=" * 60)
    print()
    
    # File path - YOUR PATH
    log_file = r"log_export.log"
    
    print(f"Loading log file: {log_file}")
    df = load_firewall_logs(log_file)
    
    print("\n" + "=" * 60)
    print("GÉNÉRATION DU RAPPORT")
    print("=" * 60)
    print()
    
    #here w save
    output_path = r'Rapport_Analyse_Firewall_IPTABLES.docx'
    create_report(df, output_path)
    
    print("\n" + "=" * 60)
    print(" ANALYSE TERMINÉE")
    print("=" * 60)

if __name__ == "__main__":
    main()